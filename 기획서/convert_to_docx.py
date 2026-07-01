from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH = "/home/user/BOOKS/기획서/AI-History_실무기획서.md"
OUT_PATH = "/home/user/BOOKS/기획서/AI-History_실무기획서.docx"

# ── 유틸리티 ──────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:color'), kwargs.get('color', '4472C4'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def bold_inline(para, text):
    """**굵게** 처리된 인라인 텍스트 처리"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

BODY_PT = 10   # 기본 본문 글자 크기

def apply_run_style(run, font_name='맑은 고딕', size=BODY_PT, color=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def set_para_spacing(para, before=0, after=4, line=240):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

# ── 스타일 설정 ───────────────────────────────────────────────
def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(BODY_PT)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 페이지 여백
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(3.0)

# ── 표 파싱·삽입 ──────────────────────────────────────────────
def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\|[-| :]+\|$', line):
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            rows.append(cells)
    return rows

def add_table(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 컬럼 너비 균등 배분
    total_width = Cm(14)
    col_w = total_width / ncols
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_w

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            # <br> 줄바꿈 처리
            parts = re.split(r'<br\s*/?>', cell_text)
            for i, part in enumerate(parts):
                if i > 0:
                    para.add_run('\n')
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', part)
                run = para.add_run(clean)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(9)
                if is_header:
                    run.bold = True
            # 표는 무채색: 음영·강조색 없이 괘선만, 헤더는 굵게만 구분
    doc.add_paragraph()

# ── 코드블록 처리 ─────────────────────────────────────────────
def add_code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    set_para_spacing(para, before=0, after=2, line=220)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    # 배경색 흉내 (단락 음영)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F6F9')
    pPr.append(shd)

# ── 헤딩 삽입 ────────────────────────────────────────────────
def add_heading(doc, text, level):
    # level 1=표지제목, 2=대장(Ⅰ~Ⅵ), 3=소절(1.2.3.), 4=세부항목
    para = doc.add_paragraph()
    set_para_spacing(para, before=120 if level <= 2 else 60,
                     after=60 if level <= 2 else 30, line=260)
    run = para.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x2E, 0x5A)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x87)
        # 하단 테두리
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '2E4A87')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x87)
    else:
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x2E, 0x5A)

# ── 본문 단락 삽입 ────────────────────────────────────────────
def add_body(doc, text, indent_level=0):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=3, line=240)
    para.paragraph_format.left_indent = Cm(0.5 * indent_level)
    bold_inline(para, text)
    for run in para.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(BODY_PT)

# ── 구분선 ───────────────────────────────────────────────────
def add_hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(para, before=60, after=60)

# ── 표지 페이지 ──────────────────────────────────────────────
def add_cover(doc, title_top, title_main, meta):
    for _ in range(6):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)

    title1 = doc.add_paragraph()
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title1.add_run(title_top)
    r.font.name = '맑은 고딕'; r.font.size = Pt(18); r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x2E, 0x5A)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    set_para_spacing(title1, before=0, after=10)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = title2.add_run(title_main)
    r2.font.name = '맑은 고딕'; r2.font.size = Pt(20); r2.bold = True
    r2.font.color.rgb = RGBColor(0x2E, 0x4A, 0x87)
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    set_para_spacing(title2, before=0, after=40)

    # 구분선
    add_hr(doc)

    for _ in range(4):
        p = doc.add_paragraph(); set_para_spacing(p, 0, 0)

    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{label}: ')
        r.font.name = '맑은 고딕'; r.font.size = Pt(11); r.bold = True
        r.font.color.rgb = RGBColor(0x2E, 0x4A, 0x87)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r2 = p.add_run(value)
        r2.font.name = '맑은 고딕'; r2.font.size = Pt(11)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        set_para_spacing(p, before=0, after=8)

    doc.add_page_break()

# ── 메인 변환 ────────────────────────────────────────────────
def convert(md_path, out_path, title_top, title_main, meta, cover_skip=10):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)
    add_cover(doc, title_top, title_main, meta)

    i = 0
    in_code = False
    code_buf = []
    table_buf = []
    in_table = False
    skip_cover = True   # 표지 메타 정보(#, ---, **작성일** 등) 건너뛰기
    cover_skip_count = 0

    while i < len(lines):
        raw = lines[i].rstrip('\n')
        stripped = raw.strip()

        # ── 표지 메타 정보 스킵 ──
        if skip_cover:
            cover_skip_count += 1
            if cover_skip_count >= cover_skip:
                skip_cover = False
            i += 1
            continue

        # ── 코드블록 ──
        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, '\n'.join(code_buf))
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── 테이블 ──
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                add_table(doc, parse_table(table_buf))
                table_buf = []
                in_table = False

        # ── 구분선 ──
        if stripped in ('---', '***', '___'):
            add_hr(doc)
            i += 1
            continue

        # ── 빈 줄 ──
        if not stripped:
            i += 1
            continue

        # ── 인용구(>) : 독자 안내·박스형 메모 ──
        if stripped.startswith('>'):
            note = stripped.lstrip('>').strip()
            if not note:
                i += 1
                continue
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            set_para_spacing(para, before=20, after=20, line=240)
            # 좌측 강조 테두리 + 연한 음영
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F5FB')
            pPr.append(shd)
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8'); left.set(qn('w:color'), '2E4A87')
            pBdr.append(left)
            pPr.append(pBdr)
            bold_inline(para, note)
            for run in para.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x87)
            i += 1
            continue

        # ── 헤딩 ──
        if stripped.startswith('#### '):
            add_heading(doc, stripped[5:], 4)
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:], 3)
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:], 2)
        elif stripped.startswith('# '):
            # 본문 내 # 헤딩은 이미 표지에서 처리, 건너뜀
            pass

        # ── □ 대항목 ──
        elif stripped.startswith('□ '):
            add_body(doc, stripped, indent_level=0)

        # ── ○ 소항목 ──
        elif stripped.startswith('○ '):
            add_body(doc, stripped, indent_level=1)

        # ── - 세부항목 ──
        elif re.match(r'^-\s+', stripped):
            add_body(doc, '  ' + stripped, indent_level=2)

        # ── * 주석/각주 ──
        elif stripped.startswith('* ') or stripped.startswith('※'):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.5)
            set_para_spacing(para, before=0, after=2, line=220)
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            run = para.add_run(clean)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.italic = True

        # ── ⅠⅡⅢ로 시작하는 파트 제목 (목차 내) ──
        elif re.match(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]', stripped):
            add_body(doc, stripped, indent_level=0)

        # ── 일반 본문 ──
        else:
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            if clean:
                add_body(doc, clean, indent_level=0)

        i += 1

    # 마지막 테이블 잔여 처리
    if in_table and table_buf:
        add_table(doc, parse_table(table_buf))

    doc.save(out_path)
    print(f"저장 완료: {out_path}")

# ── 산출물 1: 상세 실무 기획서 ──
convert(
    MD_PATH, OUT_PATH,
    'AI-History : 한일과거사 통합디지털 아카이브',
    '구축 사업 기획(안)',
    [('작성일', '2026년 6월'),
     ('작성부서', '행정안전부 사회통합지원과'),
     ('비고', '내부 기획 검토용 초안')],
)

# ── 산출물 2: 장관 보고본 (2~3쪽) ──
convert(
    "/home/user/BOOKS/기획서/AI-History_장관보고본.md",
    "/home/user/BOOKS/기획서/AI-History_장관보고본.docx",
    'AI-History : 한일과거사 통합디지털 아카이브',
    '장관 보고본',
    [('작성일', '2026년 6월'),
     ('작성부서', '행정안전부 사회통합지원과'),
     ('보고 성격', '장관 보고용 요약본')],
)
