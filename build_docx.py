# -*- coding: utf-8 -*-
"""AI 국민 라이브러리 사업기획서 (정부보고서 개조식) 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "맑은 고딕"
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x40, 0x40, 0x40)

doc = Document()

# ---------- 기본 스타일 ----------
normal = doc.styles["Normal"]
normal.font.name = KFONT
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), KFONT)

# 페이지 여백
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)


def set_font(run, size=11, bold=False, color=None, name=KFONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)


def para(text="", size=11, bold=False, color=None, align=None,
         before=0, after=4, indent=None, hanging=None, line=1.3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.left_indent = Cm(indent)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, color)
    return p


def bullet(level, text, bold=False, after=3):
    """개조식 계층: 1=□ 2=○ 3=- 4=·"""
    marks = {1: "□", 2: "○", 3: "-", 4: "·"}
    indent = {1: 0.0, 2: 0.7, 3: 1.35, 4: 2.0}
    m = marks[level]
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent[level] + 0.5)
    pf.first_line_indent = Cm(-0.5)
    pf.space_before = Pt(0); pf.space_after = Pt(after)
    pf.line_spacing = 1.3
    r = p.add_run(f"{m} ")
    set_font(r, 11, bold or level == 1, NAVY if level == 1 else None)
    r2 = p.add_run(text)
    set_font(r2, 11, bold or level == 1, NAVY if level == 1 else None)
    return p


def heading(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{num}. {title}")
    set_font(r, 15, True, NAVY)
    # 밑줄 border
    pPr = p._element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "1F3A5F")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def make_table(headers, rows, widths=None, header_fill="1F3A5F"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h); set_font(r, 10, True, RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.CENTER
            p.alignment = align
            r = p.add_run(str(val)); set_font(r, 10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================================================
# 표지
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = para("2027년도 행정안전부 신규 예산 사업(안)", size=13, bold=True,
         color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI 국민 라이브러리"); set_font(r, 32, True, NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("( AI Bravo My Life )"); set_font(r, 18, True, GRAY)

p = para("80세 이상 어르신 AI 구술채록·자서전 제작을 통한", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, before=24)
p = para("인구소멸 대응 · 과거사 기억 보존 · 청년 일자리 창출 사업", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=60)

# 표지 하단 정보 박스
info = make_table(
    ["구  분", "내  용"],
    [["소관부처", "행정안전부"],
     ["사업기간", "2027 ~ 2029년 (3개년)"],
     ["2027년 소요예산", "18,000백만원 (180억원)"],
     ["추진주체", "대한노인회 · 청년 · 사회연대경제 조직"],
     ["작 성 일", "2026. 7."]],
    widths=[5.0, 11.0], header_fill="1F3A5F")

doc.add_page_break()

# ============================================================
# 요약 (1p)
# ============================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("사 업 요 약"); set_font(r, 17, True, NAVY)
p.paragraph_format.space_after = Pt(10)

bullet(1, "사업개요")
bullet(2, "(사업명) AI 국민 라이브러리(AI Bravo My Life)")
bullet(2, "(내용) 80세 이상 어르신의 생애를 AI로 구술채록하여 개인 자서전으로 제작하고, 국가 디지털 아카이브로 영구 보존")
bullet(2, "(대상) ① 인구소멸(감소)지역 어르신 ② 과거사 피해 경험 어르신")

bullet(1, "추진배경")
bullet(2, "초고령·인구소멸 시대, 어르신의 생애 기억이 기록되지 못한 채 소멸 → 지역 정체성·국가 기억 자산의 유실")
bullet(2, "과거사 피해 고령 생존자의 증언 소멸 위기 → 기록·치유의 시급성")
bullet(2, "인구소멸지역·과거사·마을기업이 모두 행정안전부 소관 → 부처 정합성 확보")

bullet(1, "추진체계 (3주체 협업)")
bullet(2, "(대한노인회) 대상 발굴·매칭, 경로당 현장, 동행·안전 지원")
bullet(2, "(청  년) 구술채록·AI편집·디지털 아카이빙 실행 → 청년 일자리 창출")
bullet(2, "(사회연대경제 조직) 청년 교육·품질관리·권역 총괄운영")

bullet(1, "소요예산")
bullet(2, "2027년 180억원 → 2028년 350억원 → 2029년 500억원 (3개년 총 1,030억원)")

bullet(1, "기대효과")
bullet(2, "(어르신) 자존감 회복·생애 존엄 / (지역) 소멸 대응·기록 자산화")
bullet(2, "(청년) 300명 규모 신규 일자리 및 '생애기록가' 직무 창출")
bullet(2, "(국가) '국민 라이브러리' 디지털 아카이브 구축 → 세계 최초 국가 단위 AI 생애기록")

doc.add_page_break()

# ============================================================
# Ⅰ. 사업 개요
# ============================================================
heading("Ⅰ", "사업 개요")
bullet(1, "사업명")
bullet(2, "국문 : AI 국민 라이브러리")
bullet(2, "영문(브랜드) : AI Bravo My Life")
bullet(1, "사업목적")
bullet(2, "AI 기술을 활용해 어르신의 생애를 구술채록·자서전화하여 개인의 존엄을 회복하고,")
bullet(3, "이를 국가 디지털 아카이브('국민 라이브러리')로 축적하여 지역·국가의 기억 자산으로 보존")
bullet(2, "실행인력으로 청년을 활용하여 양질의 청년 일자리를 창출")
bullet(1, "사업기간 : 2027 ~ 2029년(3개년), 2027년 신규 편성")
bullet(1, "소관 및 추진")
bullet(2, "(소관) 행정안전부")
bullet(2, "(추진) 광역·기초지자체 + 사회연대경제 권역운영기관(위탁) + 대한노인회 + 청년")
bullet(1, "핵심 차별성")
bullet(2, "AI(음성인식·생성형)를 활용한 대량·저비용·고품질 생애기록의 표준화")
bullet(2, "단순 복지가 아닌 '기록 → 일자리 → 아카이브'의 선순환 구조")

# ============================================================
# Ⅱ. 추진 배경 및 필요성
# ============================================================
heading("Ⅱ", "추진 배경 및 필요성")
bullet(1, "정책 환경")
bullet(2, "(인구소멸) 행안부, 인구감소지역 89곳 지정(고시 제2021-66호) 및 관심지역 18곳 운영")
bullet(3, "지방소멸대응기금 연 1조원 투입 → 사업 재원 및 정책 연계 가능")
bullet(2, "(과거사) 진실·화해를 위한 과거사정리위원회 2기 운영, 고령 생존자·유족 다수")
bullet(3, "행정안전부 과거사관련업무지원단 소관 → 부처 정합성")
bullet(1, "문제 인식")
bullet(2, "초고령 어르신의 생애 기억이 기록되지 못한 채 매년 소멸")
bullet(2, "인구소멸지역은 '사람의 소멸'과 동시에 '기억의 소멸' 진행")
bullet(2, "과거사 피해 증언은 생존자 고령화로 확보 시한이 임박")
bullet(1, "사업 필요성")
bullet(2, "(어르신) 생애 회고를 통한 자아통합감 제고, 우울·고립 완화")
bullet(2, "(지역) 마을·지역 기억의 아카이브화 → 지역 정체성·관광·교육 자원화")
bullet(2, "(청년) 지역 기반 디지털 직무 일자리 창출 및 정주 유도")
bullet(2, "(국가) 국가 기억 자산의 체계적 축적, AI 공공활용 선도모델")
bullet(1, "선행사례 (타당성 검증)")
bullet(2, "청송군 근현대 생애사 구술기록 수집사업(마을 스토리북 제작, 2022~)")
bullet(2, "대학·노인복지관의 어르신 자서전·구술생애사 프로그램 다수 → 수요·방법론 검증")

# ============================================================
# Ⅲ. 사업 대상 [팀원 1]
# ============================================================
heading("Ⅲ", "사업 대상")
bullet(1, "대상 구조 : 2개 트랙 + 공통 자격기준")
make_table(
    ["트랙", "대상", "선정 명분", "우선 발굴지역"],
    [["A", "인구소멸지역 어르신", "지역소멸 대응·지역기억 보존", "인구감소지역 89곳 + 관심지역 18곳"],
     ["B", "과거사 피해 경험 어르신", "역사 기록·명예회복·치유", "진실화해위 신청·확정 사건 지역"]],
    widths=[1.3, 4.2, 5.0, 5.5])

bullet(1, "공통 자격기준")
bullet(2, "(연령) 만 80세 이상 원칙")
bullet(3, "트랙 B(과거사)는 사안 특성·증언 시급성 고려 만 75세 이상 탄력 적용")
bullet(2, "(우선순위) ① 초고령(85세↑) ② 건강 취약 ③ 증언 소멸위험 ④ 독거·무연고")

bullet(1, "규모 추계")
make_table(
    ["구분", "2027년(1차)", "2028년(2차)", "2029년(3차)", "누계"],
    [["트랙 A (인구소멸)", "2,100명", "3,500명", "5,000명", "10,600명"],
     ["트랙 B (과거사)", "900명", "1,500명", "2,000명", "4,400명"],
     ["합  계", "3,000명", "5,000명", "7,000명", "15,000명"]],
    widths=[4.5, 2.8, 2.8, 2.8, 3.0])

bullet(1, "대상 발굴체계")
bullet(2, "대한노인회 시군구 지회 + 읍면동 경로당(약 6.8만개) 네트워크 활용")
bullet(2, "트랙 B는 진실화해위·과거사 지원단 협조로 대상 명단 연계(개인정보 동의 전제)")
bullet(2, "지자체 사회복지·보건소 연계로 건강취약 어르신 우선 매칭")

# ============================================================
# Ⅳ. 추진 체계 [팀원 2]
# ============================================================
heading("Ⅳ", "추진 체계")
bullet(1, "거버넌스 (4단 체계)")
bullet(2, "행정안전부(총괄·예산) → 광역지자체(지역총괄) → 권역운영기관(집행·관리) → 현장(노인회·청년)")

bullet(1, "3주체 역할 분담")
make_table(
    ["주체", "역할", "핵심 기능"],
    [["대한노인회", "대상 발굴·현장", "대상 매칭, 경로당 인터뷰 공간, 동행·안전, 신뢰형성"],
     ["청년", "실행인력(일자리)", "구술채록, AI 전사·편집, 사진·기록물 디지털화, 아카이브 등록"],
     ["사회연대경제 조직", "교육·총괄운영", "청년 교육·훈련, 품질관리(QC), 정산, 권역 총괄"]],
    widths=[3.5, 3.2, 9.3])

bullet(1, "사회연대경제 총괄운영기관 (중간지원조직)")
bullet(2, "유형 : (사회적)협동조합·사회적기업·마을기업(행안부 소관)·자활기업 중 공모 선정")
bullet(2, "배치 : 전국 5대 권역(수도권·강원, 충청, 호남, 대경, 동남권)에 권역운영기관 1개소")
bullet(2, "역할 : 청년 채용·교육, 구술채록 품질관리, 저작권·윤리 관리, 실적·정산 총괄")

bullet(1, "청년 일자리 설계")
bullet(2, "직무명 : '생애기록가(Life Archivist)' → 단순 알바가 아닌 전문 직무 정체성 부여")
bullet(2, "구성 : 권역별 채록팀(청년 2~3인 + 노인회 현장코디 1인)")
bullet(2, "성장경로 : 교육 → 현장 → 우수인력 정규 전환·창업(마을기업·협동조합) 연계")

bullet(1, "AI 기술체계")
bullet(2, "구술채록 : 음성인식(STT)로 인터뷰 자동 전사")
bullet(2, "자서전 편집 : 생성형 AI 초고 작성 → 청년·전문가 윤문·검수(Human-in-the-loop)")
bullet(2, "아카이브 : 텍스트·음성·영상·사진 통합 '국민 라이브러리' 플랫폼 등록")
bullet(2, "품질·윤리 : 사실 검증(팩트체크), 초상권·저작권 귀속(어르신), 과거사 2차가해 방지 프로토콜")

# ============================================================
# Ⅴ. 세부 사업계획 [팀원 3]
# ============================================================
heading("Ⅴ", "세부 사업계획")
bullet(1, "추진 절차 (1인당 표준 프로세스)")
bullet(2, "① 대상 발굴·동의(노인회) → ② 사전조사·라포 형성 → ③ 심층 구술 인터뷰 3~4회")
bullet(2, "④ AI 전사·초고 생성 → ⑤ 청년·전문가 편집·윤문 → ⑥ 사진·기록물 편집")
bullet(2, "⑦ 자서전 제작(양장본 10부 내외) → ⑧ 증정식 → ⑨ 디지털 아카이브 등록")

bullet(1, "연차별 추진계획")
make_table(
    ["단계", "연도", "핵심 과업", "목표"],
    [["기반구축·시범", "2027", "플랫폼 구축, 권역기관 선정, 청년 교육, 시범 채록", "3,000명 / 청년 300명"],
     ["본사업", "2028", "전 권역 확대, 품질 표준화, 아카이브 고도화", "5,000명 / 청년 500명"],
     ["확산·고도화", "2029", "전국 확산, 콘텐츠 활용(전시·교육·관광)", "7,000명 / 청년 700명"]],
    widths=[3.0, 1.6, 7.4, 4.0])

bullet(1, "성과지표 (KPI)")
make_table(
    ["구분", "지표", "2027 목표"],
    [["산출", "자서전 제작 건수", "3,000건"],
     ["산출", "디지털 아카이브 등록 건수", "3,000건"],
     ["일자리", "청년 참여 인원 / 정규·창업 전환율", "300명 / 15%↑"],
     ["성과", "어르신 만족도 / 자아통합감 개선", "90%↑"],
     ["성과", "참여 지자체(시군구) 수", "30개소"]],
    widths=[2.5, 8.5, 5.0])

# ============================================================
# Ⅵ. 소요예산 [팀원 3]
# ============================================================
heading("Ⅵ", "소요예산")
bullet(1, "2027년 총사업비 : 18,000백만원 (180억원)")
make_table(
    ["세부사업", "산출근거", "예산(백만원)", "비중"],
    [["① 자서전 제작 직접비", "3,000명 × 1.2백만원(전사·인쇄·재료·플랫폼이용)", "3,600", "20.0%"],
     ["② 청년 일자리 인건비", "300명 × 2.5백만원 × 11개월(4대보험 포함)", "8,250", "45.8%"],
     ["③ 권역 총괄운영기관", "5권역 × 400백만원(운영·교육·QC)", "2,000", "11.1%"],
     ["④ 대한노인회 협력체계", "지회 30 × 30백만원 + 중앙회 총괄 200", "1,100", "6.1%"],
     ["⑤ AI 플랫폼 구축·운영", "채록–전사–편집 통합플랫폼 구축+운영", "1,500", "8.3%"],
     ["⑥ 국민 라이브러리 아카이브", "통합 디지털 아카이브(국가기록 연계)", "800", "4.5%"],
     ["⑦ 사업관리·평가·홍보", "위탁운영·성과평가·홍보·확산", "750", "4.2%"],
     ["합       계", "", "18,000", "100%"]],
    widths=[4.6, 7.4, 2.5, 1.5])

bullet(1, "재원 조달 방안")
bullet(2, "행안부 일반회계 신규사업 + 지방소멸대응기금 연계(트랙 A 대응)")
bullet(2, "지자체 대응투자(매칭) 유도로 지역 자율성·지속성 확보")

bullet(1, "3개년 중기재정 계획")
make_table(
    ["구분", "2027년", "2028년", "2029년", "3개년 계"],
    [["소요예산(억원)", "180", "350", "500", "1,030"],
     ["자서전(명)", "3,000", "5,000", "7,000", "15,000"],
     ["청년일자리(명)", "300", "500", "700", "1,500(연인원)"]],
    widths=[3.6, 2.8, 2.8, 2.8, 3.0])

# ============================================================
# Ⅶ. 기대효과 및 성과관리
# ============================================================
heading("Ⅶ", "기대효과 및 성과관리")
bullet(1, "기대효과")
bullet(2, "(사회적) 어르신 존엄 회복·세대통합, 과거사 치유·역사 정의 실현")
bullet(2, "(경제적) 청년 양질 일자리 1,500명(연인원), 지역 정주·창업 유발")
bullet(2, "(문화적) '국민 라이브러리' 국가 기억 아카이브 → 전시·교육·콘텐츠·관광 자원화")
bullet(2, "(정책적) AI 공공활용 선도, 세계 최초 국가단위 AI 생애기록 모델")
bullet(1, "성과관리")
bullet(2, "매년 성과평가(정량 KPI + 정성 만족도), 우수 권역기관 인센티브")
bullet(2, "외부 전문기관 성과평가·환류(연 1회), 3년차 종합평가 후 확대 여부 결정")
bullet(1, "위험요인 및 대응")
bullet(2, "(개인정보·윤리) 사전 동의·저작권 귀속 명문화, 과거사 2차가해 방지 교육 의무화")
bullet(2, "(품질 편차) 표준 매뉴얼·QC 체계, AI 산출물 전문가 검수 의무")
bullet(2, "(대상 접근성) 노인회 현장 네트워크로 방문·동행 채록 병행")

# ============================================================
# Ⅷ. 향후 추진일정
# ============================================================
heading("Ⅷ", "향후 추진일정")
make_table(
    ["시기", "추진 내용"],
    [["2026. 下", "사업 기획·예산 요구서 제출(행안부→기재부)"],
     ["2026.12", "국회 예산 확정"],
     ["2027. 1~3", "시행지침 수립, 권역운영기관 공모·선정"],
     ["2027. 3~5", "청년 채용·교육, AI 플랫폼 구축, 대상 발굴(노인회)"],
     ["2027. 5~11", "구술채록·자서전 제작·아카이브 등록"],
     ["2027.12", "자서전 증정식, 성과평가, 차년도 계획 수립"]],
    widths=[3.0, 13.0])

para("", after=10)
p = para("※ 본 기획서는 2027년도 행정안전부 신규 예산 사업 편성을 위한 기획(안)이며, "
         "예산 규모·단가·대상 규모는 관계기관 협의 및 예산심의 과정에서 조정될 수 있음.",
         size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, before=6)

out = "/home/user/BOOKS/AI_국민라이브러리_사업기획서.docx"
doc.save(out)
print("SAVED:", out)
