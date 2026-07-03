#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""국민통합 정책보고서 최종본(markdown) → 정부보고서 서식 Word(.docx) 변환"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "/home/user/BOOKS/reports/national_integration_report_final.md"
OUT = "/home/user/BOOKS/reports/국민통합_정책보고서.docx"
FONT = "맑은 고딕"

doc = Document()

# 기본 스타일(본문) 설정
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = color

def body(text, indent_cm, size=11, bold=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    set_font(p.add_run(text), size=size, bold=bold)
    return p

lines = open(SRC, encoding="utf-8").read().splitlines()

# 제목 페이지 요소
i = 0
in_meta = True   # 상단 메타(- 로 시작) 블록 여부
while i < len(lines):
    raw = lines[i]
    line = raw.rstrip()
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # 제목
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(12)
        set_font(p.add_run(line[2:].strip()), size=20, bold=True)
        i += 1
        continue

    # 구분선 — 최초 구분선에서 상단 메타 블록 종료
    if stripped == "---":
        in_meta = False
        i += 1
        continue

    # 상단 메타(- 로 시작하고 아직 본문 진입 전)
    if line.startswith("- ") and in_meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(line[2:].strip()), size=9, color=RGBColor(0x59, 0x59, 0x59))
        i += 1
        continue

    # ## 대제목 (Ⅰ. Ⅱ. ...)
    if line.startswith("## "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        # 하단 테두리
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "12",
            qn("w:space"): "2", qn("w:color"): "1F4E79"})
        pbdr.append(bottom)
        pPr.append(pbdr)
        set_font(p.add_run(line[3:].strip()), size=15, bold=True,
                 color=RGBColor(0x1F, 0x4E, 0x79))
        i += 1
        continue

    # ### 중제목 (1. 2. ...)
    if line.startswith("### "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(line[4:].strip()), size=12.5, bold=True,
                 color=RGBColor(0x2E, 0x2E, 0x2E))
        i += 1
        continue

    # ※ 비고
    if stripped.startswith("※"):
        body(stripped, indent_cm=0.2, size=9.5)
        i += 1
        continue

    # 개조식 계층 기호
    first = stripped[0]
    if first == "□":
        body(stripped, indent_cm=0.0, size=11.5, bold=True, space_after=3)
    elif first == "○":
        body(stripped, indent_cm=0.7, size=11)
    elif first == "-":
        body("- " + stripped[1:].strip(), indent_cm=1.5, size=10.5)
    elif first == "·":
        body("· " + stripped[1:].strip(), indent_cm=2.2, size=10.5)
    else:
        body(stripped, indent_cm=0.5, size=11)
    i += 1

doc.save(OUT)
print("saved:", OUT)
